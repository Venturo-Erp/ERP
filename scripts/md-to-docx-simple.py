#!/usr/bin/env python3
"""
簡單的 Markdown 轉 Word 工具（乾淨版本，無重複編號）
"""

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re

def create_contract_from_md(md_file, output_file):
    doc = Document()
    
    # 設定預設字體
    style = doc.styles['Normal']
    style.font.name = '新細明體'
    style.font.size = Pt(12)
    
    with open(md_file, 'r', encoding='utf-8') as f:
        lines = f.readlines()
    
    for line in lines:
        line = line.rstrip()
        
        # 標題（# 開頭）
        if line.startswith('# '):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(line[2:])
            run.font.size = Pt(18)
            run.font.bold = True
            
        # 二級標題（## 開頭）
        elif line.startswith('## '):
            p = doc.add_paragraph(line[3:])
            p.style = 'Heading 2'
            
        # 粗體（**文字**）
        elif '**' in line:
            p = doc.add_paragraph()
            parts = re.split(r'\*\*(.*?)\*\*', line)
            for i, part in enumerate(parts):
                if i % 2 == 0:
                    p.add_run(part)
                else:
                    run = p.add_run(part)
                    run.font.bold = True
                    
        # 分隔線（---）
        elif line.strip() == '---':
            doc.add_paragraph()
            
        # 表格標題行（| 開頭）
        elif line.startswith('|') and '|' in line[1:]:
            # 忽略 Markdown 表格（這裡簡化處理）
            continue
            
        # 一般段落
        elif line.strip():
            doc.add_paragraph(line)
    
    doc.save(output_file)
    print(f'✅ Word 文件已生成：{output_file}')

if __name__ == '__main__':
    import sys
    if len(sys.argv) != 3:
        print('使用方式：python3 md-to-docx-simple.py input.md output.docx')
        sys.exit(1)
    
    create_contract_from_md(sys.argv[1], sys.argv[2])
